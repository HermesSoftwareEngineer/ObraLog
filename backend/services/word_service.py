"""Word export for diários de obra using python-docx."""
from __future__ import annotations

import io
from datetime import date

from docx.shared import RGBColor  # module-level so helpers can use it


_CAMPOS_ORDER = [
    ("estaca_inicial", "Est. Inicial"),
    ("estaca_final", "Est. Final"),
    ("estaca", "Localização"),
    ("resultado", "Resultado"),
    ("lado_pista", "Lado"),
    ("tempo_manha", "Manhã"),
    ("tempo_tarde", "Tarde"),
]


def _fetch_image_bytes(img: dict, timeout: int = 6) -> bytes | None:
    url = img.get("external_url") or ""
    if url and url.startswith(("http://", "https://")):
        try:
            import urllib.request
            req = urllib.request.Request(url, headers={"User-Agent": "ObraLog/1.0"})
            with urllib.request.urlopen(req, timeout=timeout) as resp:
                return resp.read()
        except Exception:
            pass
    storage_path = img.get("storage_path") or ""
    if storage_path:
        try:
            from pathlib import Path
            p = Path(storage_path)
            if p.exists():
                return p.read_bytes()
        except Exception:
            pass
    return None


def _campo_value(r: dict, chave: str) -> str:
    if chave in ("estaca_inicial", "estaca_final", "resultado"):
        v = r.get(chave)
        return f"{v:.2f}" if isinstance(v, (int, float)) else ("—" if v is None else str(v))
    if chave == "estaca":
        return str(r.get("estaca") or "—")
    if chave == "lado_pista":
        return str(r.get("lado_pista") or "—")
    if chave in ("tempo_manha", "tempo_tarde"):
        return str(r.get(chave) or "—")
    return str((r.get("metadata_json") or {}).get(chave) or "—")


def _build_columns(schema: dict) -> list[tuple[str, str]]:
    campos_ativos = schema.get("campos_ativos") or {}
    campos_extras = schema.get("campos_extras") or []
    cols = []
    for chave, label in _CAMPOS_ORDER:
        if campos_ativos.get(chave):
            cols.append((chave, label))
    for extra in campos_extras:
        chave = extra.get("key") or extra.get("chave") or ""
        label = extra.get("label") or chave
        if chave:
            cols.append((chave, label))
    if not cols:
        cols = [("resultado", "Resultado"), ("tempo_manha", "Manhã"), ("tempo_tarde", "Tarde")]
    return cols


def _set_heading_color(paragraph, rgb: tuple[int, int, int]) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(*rgb)


def _set_cell_bg(cell, hex_color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _vmerge_cell(cell, restart: bool) -> None:
    """Mark a cell as vertical merge start (restart=True) or continuation."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vm = OxmlElement("w:vMerge")
    if restart:
        vm.set(qn("w:val"), "restart")
    tcPr.append(vm)


def gerar_word_diario(diario: dict, registros: list[dict], frentes_schemas: dict | None = None) -> bytes:
    """Generate a Word document for a diário de obra. Returns raw docx bytes."""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    frentes_schemas = frentes_schemas or {}

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    obra_nome = diario.get("obra_nome") or f"Obra {diario.get('obra_id', '')}"
    tipo = diario.get("tipo", "diario").capitalize()
    data_inicio = _fmt_date(str(diario.get("data_inicio", "")))
    data_fim = _fmt_date(str(diario.get("data_fim", "")))
    periodo = f"{data_inicio}" if data_inicio == data_fim else f"{data_inicio} a {data_fim}"
    versao = diario.get("versao_atual", 1)
    status = diario.get("status", "rascunho").capitalize()
    gerado_por = diario.get("gerado_por_nome", "sistema")
    tenant_nome = diario.get("tenant_nome", "")

    title = doc.add_heading(f"Diário de Obra — {tipo}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_heading_color(title, (45, 106, 79))

    sub = doc.add_paragraph(obra_nome)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(14)

    doc.add_paragraph()

    info_tbl = doc.add_table(rows=5, cols=2)
    info_tbl.style = "Table Grid"
    for i, (label, value) in enumerate([
        ("Período", periodo), ("Versão", str(versao)), ("Status", status),
        ("Empresa", tenant_nome), ("Gerado por", gerado_por),
    ]):
        info_tbl.cell(i, 0).text = label
        info_tbl.cell(i, 0).paragraphs[0].runs[0].bold = True
        _set_cell_bg(info_tbl.cell(i, 0), "B7E4C7")
        info_tbl.cell(i, 1).text = value

    doc.add_paragraph()

    if not registros:
        doc.add_paragraph("Nenhum registro no período.")
    else:
        grupos: dict = {}
        for r in registros:
            fid = r.get("frente_servico_id")
            grupos.setdefault(fid, []).append(r)

        for fid, regs in grupos.items():
            schema = frentes_schemas.get(fid, {})
            frente_nome = schema.get("nome") or regs[0].get("frente_servico_nome") or "Sem frente"
            cols = _build_columns(schema)

            heading = doc.add_heading(f"Frente: {frente_nome}", level=1)
            _set_heading_color(heading, (45, 106, 79))

            col_names = ["Data"] + [label for _, label in cols] + ["Observação", "Status", "Fotos"]
            n_data_cols = len(col_names) - 1  # all except Fotos

            tbl = doc.add_table(rows=1, cols=len(col_names))
            tbl.style = "Table Grid"
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            hdr_cells = tbl.rows[0].cells
            for i, name in enumerate(col_names):
                hdr_cells[i].text = name
                run = hdr_cells[i].paragraphs[0].runs[0]
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                _set_cell_bg(hdr_cells[i], "2D6A4F")

            for seq, r in enumerate(regs, start=1):
                imagens = r.get("imagens") or []
                img_data: list[bytes] = []
                for img in imagens:
                    b = _fetch_image_bytes(img)
                    if b:
                        img_data.append(b)

                n_photo_rows = max(1, len(img_data))
                bg = "F5F5F5" if seq % 2 == 0 else "FFFFFF"

                data_values = (
                    [_fmt_date(str(r.get("data", "")))]
                    + [_campo_value(r, chave) for chave, _ in cols]
                    + [str(r.get("observacao") or ""), str(r.get("status") or "—")]
                )

                first_row = tbl.add_row()
                for ci, val in enumerate(data_values):
                    first_row.cells[ci].text = val
                    _set_cell_bg(first_row.cells[ci], bg)
                    if n_photo_rows > 1:
                        _vmerge_cell(first_row.cells[ci], restart=True)

                # Add first image (or placeholder) in Fotos cell
                foto_cell = first_row.cells[-1]
                _set_cell_bg(foto_cell, bg)
                if img_data:
                    para = foto_cell.paragraphs[0]
                    run = para.add_run()
                    try:
                        run.add_picture(io.BytesIO(img_data[0]), width=Inches(1.5))
                    except Exception:
                        foto_cell.text = "[imagem]"
                else:
                    foto_cell.text = "—"

                # Additional photo rows
                for pi in range(1, n_photo_rows):
                    extra_row = tbl.add_row()
                    for ci in range(n_data_cols):
                        extra_row.cells[ci].text = ""
                        _set_cell_bg(extra_row.cells[ci], bg)
                        _vmerge_cell(extra_row.cells[ci], restart=False)
                    foto_extra = extra_row.cells[-1]
                    _set_cell_bg(foto_extra, bg)
                    para = foto_extra.paragraphs[0]
                    run = para.add_run()
                    try:
                        run.add_picture(io.BytesIO(img_data[pi]), width=Inches(1.5))
                    except Exception:
                        foto_extra.text = "[imagem]"

            doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _fmt_date(value: str) -> str:
    if not value or len(value) < 10:
        return value or ""
    try:
        d = date.fromisoformat(value[:10])
        return d.strftime("%d/%m/%Y")
    except ValueError:
        return value[:10]
